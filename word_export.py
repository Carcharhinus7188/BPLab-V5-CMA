from __future__ import annotations
from pathlib import Path
from io import BytesIO
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from calculations import final_result, avg

ROOT=Path(__file__).parent
TPL=ROOT/'templates'

def _text(v):
    if v is None:return ''
    if isinstance(v,float): return f'{v:g}'
    return str(v)

def set_cell(cell,value,size=8):
    cell.text=_text(value)
    for p in cell.paragraphs:
        p.paragraph_format.space_after=Pt(0)
        for r in p.runs:
            r.font.name='е®ӢдҪ“'; r._element.rPr.rFonts.set(qn('w:eastAsia'),'е®ӢдҪ“'); r.font.size=Pt(size)

def checked(yes,yes_text='жҳҜ',no_text='еҗҰ'):
    return f'вҳ’{yes_text}  вҳҗ{no_text}' if yes else f'вҳҗ{yes_text}  вҳ’{no_text}'

def choice(value,options):
    return '  '.join(('вҳ’' if value==o else 'вҳҗ')+str(o) for o in options)

def fill_pairs(doc, mapping):
    # Fill the cell immediately to the right of exact or near-exact label cells.
    for t in doc.tables:
        for row in t.rows:
            for i,c in enumerate(row.cells):
                label=''.join(c.text.split()).replace('пјҡ','')
                for key,val in mapping.items():
                    k=''.join(key.split()).replace('пјҡ','')
                    if label==k or (len(k)>=4 and label.startswith(k)):
                        if i+1<len(row.cells) and row.cells[i+1] is not c:
                            set_cell(row.cells[i+1],val)
                        break

def fill_check_table(table, checks):
    for row in table.rows[1:]:
        label=''.join(row.cells[0].text.split())
        matched=None
        for name,val in checks.items():
            if ''.join(name.split())[:10] in label or label[:10] in ''.join(name.split()): matched=val; break
        if matched is not None and len(row.cells)>=3: set_cell(row.cells[-2] if len(row.cells)>3 else row.cells[-1], checked(bool(matched),'з¬ҰеҗҲ','дёҚз¬ҰеҗҲ'))

def fill_anomalies(table, anomalies):
    for i,a in enumerate(anomalies[:max(0,len(table.rows)-1)],1):
        vals=[a.get('time',''),a.get('sample',''),a.get('description',''),a.get('impact',''),a.get('action',''),a.get('responsible',''),a.get('reviewer','')]
        for j,v in enumerate(vals[:len(table.rows[i].cells)]): set_cell(table.rows[i].cells[j],v)

def common_map(common, result):
    date=common.get('test_date','')
    return {'и®°еҪ•зј–еҸ·':common.get('record_no',''),'еҺҹе§Ӣи®°еҪ•зј–еҸ·':common.get('record_no',''),'жҠҘе‘Ҡзј–еҸ·':common.get('report_no',''),'е§”жүҳеҚ•дҪҚ':common.get('client',''),'е§”жүҳзј–еҸ·/д»»еҠЎеҚ•еҸ·':common.get('task_no',''),'е§”жүҳеҚ•зј–еҸ·':common.get('task_no',''),'е§”жүҳеҚ•еҸ·':common.get('task_no',''),'ж ·е“ҒеҗҚз§°':common.get('sample_name',''),'ж ·е“Ғзј–еҸ·/жү№еҸ·':common.get('sample_no',''),'ж ·е“Ғзј–еҸ·':common.get('sample_no',''),'и§„ж јеһӢеҸ·':common.get('model',''),'жқҗж–ҷеҗҚз§°':common.get('material',''),'жЈҖжөӢең°зӮ№':common.get('location',''),'жЈҖжөӢж—Ҙжңҹ':date,'зҺҜеўғжё©еәҰ':f"{common.get('temperature','')} в„ғ",'зӣёеҜ№ж№ҝеәҰ':f"{common.get('humidity','')} %RH",'жЈҖжөӢдәәе‘ҳ':common.get('operator',''),'еӨҚж ёдәәе‘ҳ':common.get('reviewer',''),'ж•°жҚ®дҝқеӯҳи·Ҝеҫ„':common.get('data_path',''),'еҺҹе§ӢеӣҫеғҸдҝқеӯҳи·Ҝеҫ„':common.get('data_path',''),'иө„ж–ҷдҝқеӯҳи·Ҝеҫ„/жЎЈжЎҲеҸ·':common.get('data_path',''),'жңҖз»Ҳз»“и®ә':result}

def export_doc(code,conf,common,setup,checks,rows,aux,anomalies,archives):
    path=TPL/conf['template']; doc=Document(path); result=final_result(rows)
    fill_pairs(doc,common_map(common,result))
    # No logo or company header is inserted; original template remains the master layout.
    if code=='rough':
        t=doc.tables; set_cell(t[0].cell(0,1),common.get('record_no'));set_cell(t[0].cell(0,3),common.get('report_no'));set_cell(t[0].cell(1,1),common.get('client'));set_cell(t[0].cell(1,3),common.get('test_date'));set_cell(t[0].cell(2,1),common.get('sample_name'));set_cell(t[0].cell(2,3),common.get('sample_no'))
        vals=[len(rows),common.get('material'),setup.get('instrument'),setup.get('calibration'),setup.get('stylus'),setup.get('standard_id'),setup.get('standard_nominal'),setup.get('standard_measured'),setup.get('level_x'),setup.get('level_y'),common.get('data_path'),f"{common.get('operator','')} / {common.get('reviewer','')}"]
        coords=[(1,0,1),(1,0,3),(1,3,1),(1,3,3),(1,4,1),(1,4,3),(1,5,1),(1,5,3),(1,6,1),(1,6,3),(1,7,1),(1,7,3)]
        for v,(ti,ri,ci) in zip(vals,coords): set_cell(t[ti].cell(ri,ci),v)
        for i,r in enumerate(rows[:6],1):
            vals=[i,r.get('sample_id'),r.get('direction'),r.get('ra1'),r.get('ra2'),r.get('ra3'),r.get('average'),choice(r.get('result'),['з¬ҰеҗҲ','дёҚз¬ҰеҗҲ']),r.get('note')]
            for j,v in enumerate(vals): set_cell(t[3].cell(i,j),v)
        set_cell(t[3].cell(7,7),choice(result,['з¬ҰеҗҲ','дёҚз¬ҰеҗҲ','йңҖеӨҚжЈҖ','д»…жҸҗдҫӣе®һжөӢз»“жһң']))
    elif code=='mc':
        t=doc.tables; set_cell(t[0].cell(0,1),common.get('record_no'));set_cell(t[0].cell(1,1),common.get('sample_name'));set_cell(t[0].cell(1,3),common.get('client'));set_cell(t[0].cell(2,1),common.get('report_no'));set_cell(t[0].cell(2,3),common.get('test_date'));set_cell(t[0].cell(3,1),f"{setup.get('machine','')} / {setup.get('software','')}")
        set_cell(t[1].cell(1,1),f"жё©еәҰпјҡ{common.get('temperature')} в„ғпјӣж№ҝеәҰпјҡ{common.get('humidity')} %RH")
        spec={1:setup.get('fixture'),2:setup.get('span'),3:setup.get('radius'),5:setup.get('parallel_block'),6:f"е№іиЎҢеәҰпјҡ{setup.get('parallelism')} mmпјӣй«ҳеәҰе·®пјҡ{setup.get('height_diff')} mm",8:setup.get('max_gap')}
        for ri,v in spec.items():
            if ri<len(t[2].rows): set_cell(t[2].cell(ri,2),v)
        set_cell(t[3].cell(1,1),setup.get('metal_batch'));set_cell(t[3].cell(2,1),f"EM={rows[0].get('em',0) if rows else 0} GPaпјӣжқҘжәҗпјҡ{setup.get('em_source','')}")
        for i,r in enumerate(rows[:6],1):
            vals=[r.get('sample_id'),r.get('width'),r.get('dm1'),r.get('dm2'),r.get('dm3'),r.get('dm_avg'),r.get('em'),r.get('k'),r.get('ffail'),r.get('tau'),r.get('crack_pos'),r.get('mode'),r.get('file_no'),choice(r.get('result'),['з¬ҰеҗҲ','дёҚз¬ҰеҗҲ'])]
            for j,v in enumerate(vals): set_cell(t[5].cell(i,j),v,7)
        set_cell(t[6].cell(1,1),f"6дёӘиҜ•ж ·дёӯз¬ҰеҗҲиҰҒжұӮж•°йҮҸпјҡ{sum(r.get('result')=='з¬ҰеҗҲ' for r in rows)} дёӘ");set_cell(t[6].cell(4,1),choice(result,['з¬ҰеҗҲ','дёҚз¬ҰеҗҲ','йңҖжҠҖжңҜиҜ„е®Ў','д»…жҸҸиҝ°з»“жһңдёҚдҪңеҲӨе®ҡ']))
        set_cell(t[7].cell(1,1),common.get('operator'));set_cell(t[7].cell(1,2),common.get('test_date'));set_cell(t[7].cell(2,1),common.get('reviewer'));set_cell(t[7].cell(2,2),common.get('test_date'))
    elif code=='xray':
        t=doc.tables; set_cell(t[0].cell(0,3),common.get('record_no'));set_cell(t[0].cell(2,3),common.get('report_no'));set_cell(t[0].cell(2,5),common.get('test_date'));set_cell(t[0].cell(3,1),choice(setup.get('sample_category'),['зүҷйҪҝж ·е“Ғ','йҮ‘еұһеҶ ','йҮ‘еұһжЎҘ','йҮ‘еұһеұҖйғЁеҸҜж‘ҳд№үйҪҝ','е…¶д»–']))
        set_cell(t[1].cell(0,1),common.get('client'));set_cell(t[1].cell(0,3),common.get('task_no'));set_cell(t[1].cell(1,1),common.get('sample_name'));set_cell(t[1].cell(1,3),common.get('sample_no'));set_cell(t[1].cell(2,1),len(rows));set_cell(t[1].cell(2,3),setup.get('sample_source'));set_cell(t[1].cell(3,3),common.get('location'));set_cell(t[1].cell(6,1),common.get('data_path'))
        set_cell(t[2].cell(1,1),f"{common.get('temperature')} в„ғ");set_cell(t[2].cell(2,1),f"{common.get('humidity')} %RH")
        devices=[setup.get('xray_id'),setup.get('plate_id'),setup.get('iqi_id'),setup.get('density_id'),setup.get('standard_density_id'),setup.get('software'),'','']
        for i,v in enumerate(devices,1): set_cell(t[3].cell(i,3),v)
        vals=[setup.get('standard_density_id'),setup.get('density_nominal'),setup.get('density_1'),setup.get('density_2'),setup.get('density_3'),round(avg(setup.get('density_1'),setup.get('density_2'),setup.get('density_3')),3),'','еҗҲж ј',common.get('operator')]
        for j,v in enumerate(vals): set_cell(t[4].cell(1,j),v)
        set_cell(t[5].cell(0,3),f"{setup.get('tube_voltage')} kV");set_cell(t[5].cell(1,1),f"{setup.get('tube_current')} mA");set_cell(t[5].cell(1,3),f"{setup.get('exposure')} ms");set_cell(t[5].cell(2,1),f"{setup.get('mas')} mAs");set_cell(t[5].cell(2,3),setup.get('focus'));set_cell(t[5].cell(3,3),f"з¬¬{setup.get('exposure_count')}ж¬Ў")
        for i,r in enumerate(rows[:10],1):
            vals=[i,r.get('sample_id'),r.get('tooth_pos'),choice('е®ҢеҘҪ',['е®ҢеҘҪ','ејӮеёё']),r.get('image_no'),'вҳ’е’¬еҗҲйқўжңқдёӢ',choice('жё…жҷ°' if r.get('image_valid')=='жңүж•Ҳ' else 'дёҚжё…',['жё…жҷ°','дёҚжё…']),choice(r.get('image_valid'),['жңүж•Ҳ','ж— ж•Ҳ']),choice(r.get('rephoto'),['еҗҰ','жҳҜ'])+(' '+r.get('note','') if r.get('note') else '')]
            for j,v in enumerate(vals): set_cell(t[6].cell(i,j),v,7)
        for i,a in enumerate(aux[:10],4):
            vals=['еӯ”еҪўеғҸиҙЁи®Ў',a.get('_label'),a.get('v1'),a.get('v2'),a.get('v3'),round(avg(a.get('v1'),a.get('v2'),a.get('v3')),2),a.get('note')]
            for j,v in enumerate(vals): set_cell(t[8].cell(i,j),v)
        for i,r in enumerate(rows[:6],1):
            vals=[r.get('sample_id'),r.get('image_no'),f"ROI-1:{r.get('roi1')}; ROI-2:{r.get('roi2')}; ROI-3:{r.get('roi3')}",'',r.get('thickness_est'),'еҗҰ',r.get('note')]
            for j,v in enumerate(vals): set_cell(t[9].cell(i,j),v,7)
        set_cell(t[11].cell(0,1),choice('жңүж•Ҳ' if all(r.get('image_valid')=='жңүж•Ҳ' for r in rows) else 'ж— ж•Ҳ',['жңүж•Ҳ','ж— ж•Ҳ']));set_cell(t[11].cell(0,3),'пјӣ'.join(r.get('thickness_est','') for r in rows if r.get('thickness_est')));set_cell(t[11].cell(1,1),choice('жңӘи§ҒжҳҺжҳҫејӮеёё' if all(r.get('anomaly')=='жңӘи§ҒжҳҺжҳҫејӮеёё' for r in rows) else 'и§ҒејӮеёё',['жңӘи§ҒжҳҺжҳҫејӮеёё','и§ҒејӮеёё']));set_cell(t[11].cell(1,3),choice(result,['еҗҲж ј','дёҚеҗҲж ј','йңҖеӨҚжЈҖ','и¶…еҮәж–№жі•йҖӮз”ЁиҢғеӣҙ']))
    elif code=='warp':
        t=doc.tables; set_cell(t[0].cell(0,1),common.get('record_no'));set_cell(t[0].cell(1,1),common.get('client'));set_cell(t[0].cell(1,3),common.get('report_no'));set_cell(t[0].cell(2,1),common.get('sample_name'));set_cell(t[0].cell(2,3),common.get('sample_no'))
        set_cell(t[1].cell(0,1),setup.get('receive_no'));set_cell(t[1].cell(1,1),len(rows));set_cell(t[1].cell(4,1),f"{common.get('temperature')} в„ғ");set_cell(t[1].cell(5,1),f"{common.get('humidity')} %RH");set_cell(t[1].cell(6,1),common.get('test_date'));set_cell(t[1].cell(7,1),f"ејҖе§Ӣпјҡ{common.get('start_time')} з»“жқҹпјҡ{common.get('end_time')}")
        setup_rows=[setup.get('image_machine')+'пјӣжңүж•ҲжңҹиҮіпјҡ'+str(setup.get('image_calibration','')),setup.get('software'),setup.get('cut_machine'),setup.get('fixture'),setup.get('blade'),common.get('data_path')]
        for i,v in enumerate(setup_rows,1): set_cell(t[2].cell(i,3),v)
        for i,r in enumerate(rows[:10],1):
            vals=[i,r.get('sample_id'),r.get('pre_image'),'вҳ’жҳҜ вҳҗеҗҰ','вҳ’жҳҜ вҳҗеҗҰ',r.get('h1'),r.get('operator') or common.get('operator')]
            for j,v in enumerate(vals): set_cell(t[4].cell(i,j),v,7)
            vals2=[i,r.get('sample_id'),f"{r.get('cut_start')} / {r.get('cut_end')}",choice(r.get('coolant'),['жҳҜ','еҗҰ']),choice(r.get('cut_quality'),['еҗҲж ј','дёҚеҗҲж ј']),choice(r.get('remade'),['жҳҜ','еҗҰ']),r.get('note')]
            for j,v in enumerate(vals2): set_cell(t[6].cell(i,j),v,7)
            vals3=[i,r.get('sample_id'),r.get('post_image'),'вҳ’жҳҜ вҳҗеҗҰ','вҳ’жҳҜ вҳҗеҗҰ',r.get('h2'),r.get('operator') or common.get('operator')]
            for j,v in enumerate(vals3): set_cell(t[8].cell(i,j),v,7)
            vals4=[i,r.get('sample_id'),r.get('h1'),r.get('h2'),r.get('delta'),f"Вұ{setup.get('limit')} mm" if float(setup.get('limit') or 0)>0 else 'д»…и®°еҪ•',choice(r.get('result'),['еҗҲж ј','дёҚеҗҲж ј','д»…и®°еҪ•']),r.get('note')]
            for j,v in enumerate(vals4): set_cell(t[9].cell(i,j),v,7)
        set_cell(t[13].cell(0,1),common.get('operator'));set_cell(t[13].cell(0,3),common.get('test_date'));set_cell(t[13].cell(1,1),common.get('reviewer'));set_cell(t[13].cell(1,3),common.get('test_date'))
    elif code=='cte':
        t=doc.tables; fill_pairs(doc,common_map(common,result)); set_cell(t[0].cell(0,1),common.get('record_no'));set_cell(t[0].cell(0,3),common.get('report_no'));set_cell(t[0].cell(2,3),common.get('test_date'))
        set_cell(t[1].cell(0,1),common.get('client'));set_cell(t[1].cell(0,3),common.get('task_no'));set_cell(t[1].cell(1,1),common.get('sample_name'));set_cell(t[1].cell(1,3),common.get('sample_no'));set_cell(t[1].cell(2,1),common.get('material'));set_cell(t[1].cell(2,3),len(rows));set_cell(t[1].cell(6,1),common.get('data_path'))
        for i,r in enumerate(rows[:6],1):
            vals=[i,r.get('sample_id'),r.get('material'),r.get('l0'),r.get('width'),r.get('thickness'),'жӯЈзЎ®',r.get('pv'),choice('жҳҜ',['жҳҜ','еҗҰ']),r.get('note')]
            for j,v in enumerate(vals): set_cell(t[5].cell(i,j),v,7)
            vals2=[i,r.get('sample_id'),r.get('start_time'),r.get('end_time'),r.get('start_temp'),r.get('end_temp'),choice('жӯЈеёё',['жӯЈеёё','ејӮеёё']),choice('жҳҜ',['жҳҜ','еҗҰ']),r.get('curve_file'),r.get('report_file'),choice(r.get('valid'),['жңүж•Ҳ','ж— ж•Ҳ'])]
            for j,v in enumerate(vals2): set_cell(t[6].cell(i,j),v,7)
            vals3=[i,r.get('sample_id'),f"{r.get('start_temp')}пҪһ{r.get('end_temp')}",r.get('l0'),r.get('dt'),r.get('dl_um'),r.get('alpha'),setup.get('limit'),choice(r.get('result'),['еҗҲж ј','дёҚеҗҲж ј','д»…и®°еҪ•','йңҖеӨҚжөӢ']),r.get('note')]
            for j,v in enumerate(vals3): set_cell(t[8].cell(i,j),v,7)
        for i,a in enumerate(aux[:13],1):
            vals=[i,a.get('temp') or a.get('_label'),a.get('dl_um'),round(float(a.get('dl_um') or 0)/1000,6),'','','вҳ’'+a.get('source','иҪҜд»¶'),a.get('note')]
            for j,v in enumerate(vals): set_cell(t[7].cell(i,j),v,7)
        alphas=[r.get('alpha',0) for r in rows if r.get('valid')=='жңүж•Ҳ'];set_cell(t[11].cell(1,1),round(sum(alphas)/len(alphas),3) if alphas else 0);set_cell(t[11].cell(0,1),choice(result,['еҗҲж ј','дёҚеҗҲж ј','д»…жҸҗдҫӣе®һжөӢеҖј','йңҖеӨҚжөӢ']))
    elif code=='shock':
        t=doc.tables; fill_pairs(doc,common_map(common,result));set_cell(t[0].cell(0,1),common.get('client'));set_cell(t[0].cell(0,3),common.get('task_no'));set_cell(t[0].cell(1,1),common.get('report_no'));set_cell(t[0].cell(1,3),common.get('record_no'));set_cell(t[0].cell(2,1),common.get('sample_name'));set_cell(t[0].cell(2,3),common.get('sample_no'));set_cell(t[0].cell(4,1),len(rows));set_cell(t[0].cell(4,3),common.get('test_date'))
        set_cell(t[1].cell(1,2),f"{common.get('temperature')} в„ғ");set_cell(t[1].cell(2,2),f"{common.get('humidity')} %RH");set_cell(t[1].cell(3,2),f"{setup.get('illumination')} lx")
        devices=[setup.get('oven'),setup.get('thermometer'),setup.get('stopwatch'),setup.get('magnifier'),setup.get('light'),setup.get('perforated_container'),setup.get('ice_container')]
        for i,v in enumerate(devices,1): set_cell(t[2].cell(i,2),v)
        values=[choice('ж— ејӮеёё',['ж— ејӮеёё','жңүејӮеёё']),f"и®ҫе®ҡ/зЁіе®ҡпјҡ{setup.get('oven_temp')}в„ғ",f"ж—¶й•ҝпјҡ{setup.get('first_heat')}min",setup.get('soap'),f"{setup.get('ice_temp')}в„ғ",f"{setup.get('transfer')}s",f"{setup.get('soak')}min",f"жё©еәҰпјҡ{setup.get('oven_temp')}в„ғпјӣж—¶й•ҝпјҡ{setup.get('second_heat')}min"]
        for i,v in enumerate(values,1): set_cell(t[3].cell(i,2),v)
        for i,a in enumerate(aux[:8],1):
            vals=[i,a.get('_label'),a.get('time'),a.get('temperature'),choice(a.get('stable'),['жҳҜ','еҗҰ']),choice(a.get('status'),['з¬ҰеҗҲ','еҒҸзҰ»']),a.get('operator')]
            for j,v in enumerate(vals): set_cell(t[4].cell(i,j),v,7)
        for idx,r in enumerate(rows[:28]):
            ti=7 if idx<14 else 8; ri=idx+1 if idx<14 else idx-13
            vals=[idx+1,r.get('sample_id'),choice(r.get('initial'),['ж— ','жңү']),choice(r.get('crack'),['ж— ','жңү']),choice(r.get('chip'),['ж— ','жңү']),choice(r.get('fracture'),['ж— ','жңү']),choice(r.get('result'),['еҗҲж ј','дёҚеҗҲж ј']),r.get('description')]
            for j,v in enumerate(vals): set_cell(t[ti].cell(ri,j),v,7)
        set_cell(t[9].cell(0,1),len(rows));set_cell(t[9].cell(0,3),sum(r.get('result') in ('з¬ҰеҗҲ','дёҚз¬ҰеҗҲ') for r in rows));set_cell(t[9].cell(1,1),sum(r.get('crack')=='жңү' for r in rows));set_cell(t[9].cell(1,3),sum(r.get('chip')=='жңү' for r in rows));set_cell(t[9].cell(2,1),sum(r.get('fracture')=='жңү' for r in rows));set_cell(t[9].cell(4,1),choice(result,['еҗҲж ј','дёҚеҗҲж ј']))
    elif code=='bend':
        t=doc.tables;set_cell(t[0].cell(0,1),common.get('record_no'));set_cell(t[0].cell(0,3),common.get('report_no'));set_cell(t[0].cell(0,5),common.get('test_date'));set_cell(t[0].cell(1,1),common.get('client'));set_cell(t[0].cell(1,3),common.get('sample_name'));set_cell(t[0].cell(1,5),common.get('sample_no'));set_cell(t[0].cell(2,1),common.get('material'));set_cell(t[0].cell(2,3),setup.get('printing'));set_cell(t[0].cell(2,5),setup.get('heat_record'));set_cell(t[0].cell(3,1),choice(setup.get('direction'),['й•ҝиҪҙе№іиЎҢzиҪҙ','й•ҝиҪҙеһӮзӣҙzиҪҙпјҲx/yиҪҙпјү']))
        dev=[setup.get('machine'),setup.get('sensor'),setup.get('deflectometer'),setup.get('software'),setup.get('dimension_tool'),f"жё©еәҰ{common.get('temperature')}в„ғпјӣж№ҝеәҰ{common.get('humidity')}%RH"]
        for i,v in enumerate(dev,1):set_cell(t[1].cell(i,1),v)
        set_cell(t[2].cell(3,2),f"ж ЎеҮҶеҖј{setup.get('cal_value')}Nпјӣзі»ж•°{setup.get('cal_factor')}");set_cell(t[3].cell(1,3),setup.get('span'));set_cell(t[3].cell(3,5),common.get('data_path'));set_cell(t[3].cell(4,1),setup.get('fixture'));set_cell(t[3].cell(6,3),setup.get('parallel_block'));set_cell(t[3].cell(6,5),setup.get('gap'));set_cell(t[3].cell(7,5),setup.get('zero_force'))
        for i,r in enumerate(rows[:6],1):
            vals=[i,r.get('sample_id'),r.get('length'),r.get('width'),r.get('height'),r.get('span'),r.get('speed'),r.get('fmax'),r.get('stress'),r.get('file_no'),choice(r.get('state'),['е®Ңж•ҙ','ж–ӯиЈӮ','ејӮеёё']),choice(r.get('result'),['з¬ҰеҗҲ','дёҚз¬ҰеҗҲ']),r.get('note')]
            for j,v in enumerate(vals): set_cell(t[4].cell(i,j),v,7)
        set_cell(t[4].cell(7,1),choice(result,['е…ЁйғЁз¬ҰеҗҲ','еӯҳеңЁдёҚз¬ҰеҗҲ','йңҖеӨҚжөӢ/жҠҖжңҜиҜ„е®Ў']));set_cell(t[7].cell(3,1),f"{common.get('operator')} / {common.get('test_date')}");set_cell(t[7].cell(3,3),f"{common.get('reviewer')} / {common.get('test_date')}");set_cell(t[7].cell(4,3),choice(result,['еҗҲж ј','дёҚеҗҲж ј','д»…жҸҸиҝ°з»“жһң','йңҖеӨҚжөӢ']))
    elif code=='hv':
        t=doc.tables;set_cell(t[0].cell(0,1),common.get('record_no'));set_cell(t[0].cell(0,3),common.get('test_date'));fill_pairs(doc,common_map(common,result));set_cell(t[1].cell(0,1),common.get('temperature'));set_cell(t[1].cell(0,3),common.get('humidity'));set_cell(t[1].cell(1,1),setup.get('hardness_model'));set_cell(t[1].cell(1,3),setup.get('hardness_id'));set_cell(t[1].cell(1,5),setup.get('calibration'));set_cell(t[1].cell(3,1),setup.get('block_id'));set_cell(t[1].cell(3,3),setup.get('block_nominal'));set_cell(t[1].cell(3,5),setup.get('block_valid'));set_cell(t[1].cell(4,1),setup.get('block_allow'));set_cell(t[1].cell(4,3),setup.get('block_1'));set_cell(t[1].cell(4,5),setup.get('block_2'));set_cell(t[1].cell(5,1),setup.get('block_3'));set_cell(t[1].cell(5,3),round(avg(setup.get('block_1'),setup.get('block_2'),setup.get('block_3')),1));set_cell(t[1].cell(6,1),setup.get('surface_method'));set_cell(t[1].cell(6,3),setup.get('surface_ra'))
        for i,r in enumerate(rows[:12],1):
            vals=[r.get('sample_id'),r.get('face'),r.get('hv1'),r.get('hv2'),r.get('hv3'),r.get('average'),r.get('ra'),setup.get('force'),setup.get('dwell'),r.get('image_no'),r.get('note')]
            for j,v in enumerate(vals): set_cell(t[4].cell(i,j),v,7)
        # Summary by sample
        for s in range(1,7):
            rr=[r for r in rows if r.get('sample_id')==f'S{s}']; m1=rr[0].get('average',0) if len(rr)>0 else 0; m2=rr[1].get('average',0) if len(rr)>1 else 0; total=round(avg(m1,m2),1)
            vals=[f'S{s}',m1,m2,total,'д»…жҸҸиҝ°',setup.get('limit'),choice('дёҚеҲӨе®ҡ',['з¬ҰеҗҲ','дёҚз¬ҰеҗҲ','дёҚеҲӨе®ҡ']),'','']
            for j,v in enumerate(vals):set_cell(t[6].cell(s,j),v,7)
        set_cell(t[9].cell(0,1),common.get('operator'));set_cell(t[9].cell(0,3),common.get('test_date'));set_cell(t[9].cell(0,5),common.get('reviewer'));set_cell(t[9].cell(1,1),common.get('approver'));set_cell(t[9].cell(1,3),common.get('test_date'));set_cell(t[9].cell(1,5),choice('жңүж•Ҳ',['жңүж•Ҳ','ж— ж•Ҳ','йңҖжҠҖжңҜиҜ„е®Ў']))
    elif code=='thickness':
        t=doc.tables;set_cell(t[0].cell(0,1),common.get('record_no'));set_cell(t[0].cell(0,3),common.get('report_no'));set_cell(t[0].cell(0,5),common.get('test_date'));set_cell(t[0].cell(1,1),common.get('client'));set_cell(t[0].cell(1,3),common.get('sample_name'));set_cell(t[0].cell(1,5),common.get('sample_no'));set_cell(t[0].cell(2,1),common.get('material'));set_cell(t[0].cell(2,3),setup.get('design'));set_cell(t[0].cell(3,3),common.get('location'));set_cell(t[0].cell(3,5),f"{common.get('operator')} / {common.get('reviewer')}")
        vals=[common.get('temperature'),common.get('humidity'),f"{setup.get('balance_start')}пҪһ{setup.get('balance_end')}",setup.get('image_machine'),setup.get('calibration'),setup.get('magnification'),setup.get('block_id'),setup.get('block_nominal'),setup.get('block_measured'),round(float(setup.get('block_measured') or 0)-float(setup.get('block_nominal') or 0),4),setup.get('software'),f"{setup.get('preheat_start')}пҪһ{setup.get('preheat_end')}",setup.get('bottom_light'),common.get('data_path')]
        coords=[(1,0,1),(1,0,3),(1,0,5),(1,1,1),(1,1,3),(1,1,5),(1,2,1),(1,2,3),(1,2,5),(1,3,1),(1,3,5),(1,4,1),(1,4,3),(1,4,5)]
        for v,(ti,ri,ci) in zip(vals,coords):set_cell(t[ti].cell(ri,ci),v)
        for i,r in enumerate(rows[:5],1):
            vals=[r.get('sample_id'),r.get('f1'),r.get('f2'),r.get('f3'),r.get('m1'),r.get('m2'),r.get('m3'),r.get('e1'),r.get('e2'),r.get('e3'),r.get('fixed_avg'),r.get('mid_avg'),r.get('end_avg'),r.get('total_avg'),r.get('deviation'),choice(r.get('result'),['з¬ҰеҗҲ','дёҚз¬ҰеҗҲ','д»…и®°еҪ•'])]
            for j,v in enumerate(vals):set_cell(t[2].cell(i,j),v,7)
        set_cell(t[3].cell(1,1),choice(result,['з¬ҰеҗҲ','дёҚз¬ҰеҗҲ','д»…жҸҗдҫӣе®һжөӢеҖј','йңҖеӨҚжөӢ']));set_cell(t[3].cell(2,1),f"{common.get('operator')} / {common.get('test_date')}");set_cell(t[3].cell(2,3),f"{common.get('reviewer')} / {common.get('test_date')}")
    elif code=='color':
        t=doc.tables;fill_pairs(doc,common_map(common,result));set_cell(t[0].cell(0,1),common.get('record_no'));set_cell(t[0].cell(0,3),common.get('report_no'));set_cell(t[0].cell(1,1),common.get('client'));set_cell(t[0].cell(1,3),common.get('test_date'));set_cell(t[1].cell(0,1),common.get('sample_name'));set_cell(t[1].cell(0,3),common.get('model'));set_cell(t[1].cell(1,1),common.get('sample_no'));set_cell(t[2].cell(0,1),common.get('temperature'));set_cell(t[2].cell(0,3),common.get('humidity'))
        for i,obs in enumerate([setup.get('observer1'),setup.get('observer2'),setup.get('observer3')],1):set_cell(t[3].cell(i,1),obs)
        set_cell(t[4].cell(1,2),setup.get('instrument'));set_cell(t[5].cell(0,1),setup.get('xenon'));set_cell(t[5].cell(1,1),setup.get('filter'));set_cell(t[6].cell(1,3),setup.get('illuminance'));set_cell(t[7].cell(1,2),'ж°ҷзҒҜ');set_cell(t[7].cell(2,2),setup.get('illuminance'));set_cell(t[7].cell(3,2),setup.get('water_temp'));set_cell(t[9].cell(1,1),setup.get('exposure_hours'))
        for i,r in enumerate(rows[:12],1):
            vals=[i,r.get('sample_id'),r.get('control_id'),choice(r.get('shape'),['еңҶзүҮ','зүҷеҪў','е…¶д»–']),r.get('size'),choice(r.get('cover'),['иҜ•ж ·еӨ№','й”Ўз®”','й“қз®”','е…¶д»–']),'дёҖеҚҠ',choice('жҳҜ',['жҳҜ','еҗҰ']),r.get('position'),choice(r.get('photo'),['жҳҜ','еҗҰ']),r.get('note')]
            for j,v in enumerate(vals):set_cell(t[8].cell(i,j),v,7)
            vals2=[r.get('sample_id'),r.get('obs1'),r.get('obs2'),r.get('obs3'),choice(r.get('majority'),['жңӘи§ҒжҳҺжҳҫиүІжіҪе·®ејӮ','еҸҜи§ҒиүІжіҪе·®ејӮ','ж— жі•еҲӨе®ҡ']),'жҳҜ',setup.get('limit'),choice(r.get('result'),['еҗҲж ј','дёҚеҗҲж ј','йңҖеӨҚжЈҖ','д»…жҸҸиҝ°']),r.get('note')]
            for j,v in enumerate(vals2):set_cell(t[13].cell(i,j),v,7)
        for i,a in enumerate(aux[:7],1):
            vals=[a.get('_label'),a.get('datetime'),a.get('hours'),a.get('water_temp'),a.get('illuminance'),a.get('distance'),choice(a.get('state'),['жӯЈеёё','ејӮеёё']),choice(a.get('state'),['жӯЈеёё','ејӮеёё']),a.get('operator'),'']
            for j,v in enumerate(vals):set_cell(t[10].cell(i,j),v,7)
        set_cell(t[14].cell(0,1),choice('жңӘи§ҒжҳҺжҳҫиүІжіҪе·®ејӮ' if result=='з¬ҰеҗҲ' else ('ж— жі•еҲӨе®ҡ' if 'еӨҚжЈҖ' in result else 'еҸҜи§ҒжҳҺжҳҫиүІжіҪе·®ејӮ'),['жңӘи§ҒжҳҺжҳҫиүІжіҪе·®ејӮ','еҸҜи§ҒиҪ»еҫ®иүІжіҪе·®ејӮ','еҸҜи§ҒжҳҺжҳҫиүІжіҪе·®ејӮ','ж— жі•еҲӨе®ҡ']));set_cell(t[14].cell(0,3),choice(result,['еҗҲж ј','дёҚеҗҲж ј','йңҖеӨҚжЈҖ','дёҚдҪңз¬ҰеҗҲжҖ§еҲӨе®ҡ']))
    elif code=='cut':
        t=doc.tables;set_cell(t[0].cell(0,1),common.get('record_no'));set_cell(t[0].cell(0,3),setup.get('sop_version'));set_cell(t[0].cell(1,1),common.get('task_no'));set_cell(t[0].cell(1,3),common.get('sample_no'));set_cell(t[0].cell(2,1),common.get('client'));set_cell(t[0].cell(2,3),common.get('test_date'));set_cell(t[0].cell(3,1),common.get('operator'));set_cell(t[0].cell(3,3),common.get('reviewer'))
        for i,r in enumerate(rows[:8],1):
            vals=[i,r.get('sample_id'),r.get('name_material'),r.get('size_shape'),r.get('position'),'вҳ’дёҚеҫ—жҚҹдјӨеҫ…жөӢйқў/вҳ’дҝқз•ҷеҹәеҮҶйқў','вҳ’е®ҢеҘҪ',r.get('note','')]
            for j,v in enumerate(vals):set_cell(t[1].cell(i,j),v,7)
        setupvals=[f"и®ҫеӨҮеһӢеҸ·/зј–еҸ·пјҡ{setup.get('machine')}пјӣзӮ№жЈҖпјҡ{setup.get('maintenance')}",setup.get('blade_spec'),setup.get('fixture'),setup.get('coolant'),setup.get('safety')]
        for i,v in enumerate(setupvals,1):set_cell(t[2].cell(i,2),v)
        for i,r in enumerate(rows[:10],1):
            vals=[i,r.get('sample_id'),r.get('blade'),r.get('fixture'),choice('е·ІеҜ№еҮҶдёӯеҝғ',['е·ІеҜ№еҮҶдёӯеҝғ','еһӮзӣҙеҸ—еҠӣ']),r.get('stroke'),r.get('feed'),choice(r.get('coolant'),['иҝһз»ӯ','ејӮеёё']),r.get('start'),r.get('end'),choice(r.get('slow_feed'),['жҳҜ','еҗҰ']),choice(r.get('running'),['жӯЈеёё','ејӮе“Қ','жҢҜеҠЁ','жҖҘеҒң']),choice(r.get('quality'),['еҗҲж ј','еҙ©иҫ№','иЈӮзә№','жҚҹдјӨеҫ…жөӢйқў','йҮҚеҲ¶ж ·']),r.get('photo'),r.get('operator')]
            for j,v in enumerate(vals):set_cell(t[5].cell(i,j),v,7)
        set_cell(t[9].cell(1,1),choice('еҲҮеүІиҝҮзЁӢеҸ—жҺ§пјҢиҜ•ж ·еҸҜиҝӣе…ҘеҗҺз»ӯйҮ‘зӣёеҲ¶ж ·/жЈҖжөӢгҖӮ' if result=='з¬ҰеҗҲ' else 'еӯҳеңЁејӮеёёпјҢйңҖйҮҚж–°еҲҮеүІ/йҮҚж–°еҲ¶ж ·гҖӮ',['еҲҮеүІиҝҮзЁӢеҸ—жҺ§пјҢиҜ•ж ·еҸҜиҝӣе…ҘеҗҺз»ӯйҮ‘зӣёеҲ¶ж ·/жЈҖжөӢгҖӮ','еӯҳеңЁејӮеёёдҪҶе·ІеӨ„зҗҶпјҢдёҚеҪұе“ҚеҗҺз»ӯдҪҝз”ЁгҖӮ','еӯҳеңЁејӮеёёпјҢйңҖйҮҚж–°еҲҮеүІ/йҮҚж–°еҲ¶ж ·гҖӮ']));set_cell(t[9].cell(4,1),f"зӯҫеӯ—пјҡ{common.get('operator')} ж—Ҙжңҹпјҡ{common.get('test_date')}");set_cell(t[9].cell(5,1),f"зӯҫеӯ—пјҡ{common.get('reviewer')} ж—Ҙжңҹпјҡ{common.get('test_date')}")
    # Generic anomalies and archives are retained in the JSON draft; templates' major fields are filled above.
    out=BytesIO();doc.save(out);out.seek(0);return out,result
